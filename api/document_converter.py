#!/usr/bin/env python3
"""
Document Format Converter for PII Masking Pipeline

This module provides functionality to convert various document formats 
(TXT, DOCX) to PDF format to enable PII masking using the existing 
PDF-based masking pipeline.
"""

import os
import sys
import logging
from pathlib import Path
from typing import Optional
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.units import inch

# Configure logging
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

try:
    from docx import Document
except ImportError:
    logger.warning("python-docx not installed. Word document conversion will not work.")
    Document = None

class DocumentConverter:
    """Handles conversion of various document formats to PDF."""
    
    @staticmethod
    def txt_to_pdf(txt_path: str, output_pdf_path: str) -> bool:
        """
        Convert a text file to PDF format.
        
        Args:
            txt_path: Path to the input text file
            output_pdf_path: Path for the output PDF file
            
        Returns:
            bool: True if conversion successful, False otherwise
        """
        try:
            # Read the text file with various encoding attempts
            text_content = ""
            encodings = ['utf-8', 'latin1', 'cp1252', 'utf-16']
            
            for encoding in encodings:
                try:
                    with open(txt_path, 'r', encoding=encoding) as file:
                        text_content = file.read()
                    break
                except UnicodeDecodeError:
                    continue
            
            if not text_content:
                logger.error("Failed to read text file with any encoding")
                return False
            
            # Create PDF
            doc = SimpleDocTemplate(output_pdf_path, pagesize=letter, 
                                  rightMargin=72, leftMargin=72,
                                  topMargin=72, bottomMargin=18)
            
            # Get styles
            styles = getSampleStyleSheet()
            normal_style = styles['Normal']
            
            # Build content
            story = []
            
            # Split text into paragraphs
            paragraphs = text_content.split('\n\n')
            
            for para_text in paragraphs:
                if para_text.strip():
                    # Handle line breaks within paragraphs
                    para_text = para_text.replace('\n', '<br/>')
                    paragraph = Paragraph(para_text, normal_style)
                    story.append(paragraph)
                    story.append(Spacer(1, 12))
            
            # Build PDF
            doc.build(story)
            
            logger.info(f"Successfully converted {txt_path} to {output_pdf_path}")
            return True
            
        except Exception as e:
            logger.error(f"Error converting TXT to PDF: {e}")
            return False
    
    @staticmethod
    def docx_to_pdf(docx_path: str, output_pdf_path: str) -> bool:
        """
        Convert a Word document to PDF format.
        
        Args:
            docx_path: Path to the input DOCX file
            output_pdf_path: Path for the output PDF file
            
        Returns:
            bool: True if conversion successful, False otherwise
        """
        if Document is None:
            logger.error("python-docx not installed. Cannot convert DOCX files.")
            return False
            
        try:
            # Read the DOCX file
            doc = Document(docx_path)
            
            # Extract text content
            full_text = []
            
            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    full_text.append(paragraph.text)
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text)
                    if row_text:
                        full_text.append(' | '.join(row_text))
            
            # Create PDF
            pdf_doc = SimpleDocTemplate(output_pdf_path, pagesize=letter, 
                                      rightMargin=72, leftMargin=72,
                                      topMargin=72, bottomMargin=18)
            
            # Get styles
            styles = getSampleStyleSheet()
            normal_style = styles['Normal']
            
            # Build content
            story = []
            
            for text in full_text:
                if text.strip():
                    # Escape special characters for ReportLab
                    text = text.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
                    paragraph = Paragraph(text, normal_style)
                    story.append(paragraph)
                    story.append(Spacer(1, 12))
            
            # Build PDF
            pdf_doc.build(story)
            
            logger.info(f"Successfully converted {docx_path} to {output_pdf_path}")
            return True
            
        except Exception as e:
            logger.error(f"Error converting DOCX to PDF: {e}")
            return False
    
    @staticmethod
    def convert_to_pdf(input_path: str, output_pdf_path: str) -> bool:
        """
        Convert any supported document format to PDF.
        
        Args:
            input_path: Path to the input document
            output_pdf_path: Path for the output PDF file
            
        Returns:
            bool: True if conversion successful, False otherwise
        """
        file_extension = os.path.splitext(input_path)[1].lower()
        
        if file_extension == '.txt':
            return DocumentConverter.txt_to_pdf(input_path, output_pdf_path)
        elif file_extension == '.docx':
            return DocumentConverter.docx_to_pdf(input_path, output_pdf_path)
        elif file_extension == '.pdf':
            # Already a PDF, just copy it
            import shutil
            shutil.copy2(input_path, output_pdf_path)
            logger.info(f"Copied PDF from {input_path} to {output_pdf_path}")
            return True
        else:
            logger.error(f"Unsupported file format: {file_extension}")
            return False

def main():
    """Command-line interface for document conversion."""
    if len(sys.argv) != 3:
        print("Usage: python document_converter.py <input_document> <output.pdf>")
        print("\nSupported input formats: .txt, .docx, .pdf")
        print("\nExamples:")
        print("  python document_converter.py document.txt output.pdf")
        print("  python document_converter.py document.docx output.pdf")
        sys.exit(1)
    
    input_path = sys.argv[1]
    output_path = sys.argv[2]
    
    if not os.path.exists(input_path):
        print(f"Error: Input file '{input_path}' not found")
        sys.exit(1)
    
    if not output_path.lower().endswith('.pdf'):
        print("Error: Output file must have .pdf extension")
        sys.exit(1)
    
    converter = DocumentConverter()
    success = converter.convert_to_pdf(input_path, output_path)
    
    if success:
        print(f"✓ Successfully converted {input_path} to {output_path}")
        sys.exit(0)
    else:
        print(f"✗ Failed to convert {input_path}")
        sys.exit(1)

if __name__ == "__main__":
    main()
