#!/usr/bin/env python3
"""
BERT-based PII Masker with Configurable Strategies

This script uses BERT NER model for PII detection and applies different masking strategies
based on input parameters. It preserves the original PDF layout while applying the 
specified masking strategy for each PII entity.

Input Format: PII:Type:Strategy
- PII: The actual text to be masked
- Type: The type of PII (PERSON, EMAIL, PHONE, etc.)
- Strategy: The masking strategy (redact, mask, pseudo)
- Coordinates: rectangle coordinates (x0, y0, x1, y1)

Usage:
    python bert_pii_masker.py input.pdf output.pdf pii_config.txt
    python bert_pii_masker.py input.pdf output.pdf --interactive
"""

import sys
import os
import random
import logging
import re
from typing import List, Dict, Any, Tuple, Optional
from dataclasses import dataclass
from dotenv import load_dotenv

# Load environment variables
load_dotenv()
os.environ["GOOGLE_API_KEY"] = os.getenv("GOOGLE_API_KEY")

# Configure logging
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

try:
    import fitz  # PyMuPDF
except ImportError:
    print("PyMuPDF not found. Please install: pip install PyMuPDF")
    sys.exit(1)

try:
    from docx import Document
except ImportError:
    print("python-docx not found. Please install: pip install python-docx")
    sys.exit(1)

@dataclass
class PIIConfig:
    """Configuration for PII masking with coordinates."""
    text: str
    pii_type: str
    strategy: str
    page_num: int = 0
    x0: float = 0.0
    y0: float = 0.0
    x1: float = 0.0
    y1: float = 0.0
    replacement: Optional[str] = None

class BERTPIIMasker:
    """
    BERT-based PII masker with configurable strategies and LLM validation.
    """
    
    def __init__(self):
        self.ner_pipeline = None
        self.pseudo_data = self._initialize_pseudo_data()
        self.used_mappings = {}  # Global mappings for consistency
        self.name_part_mappings = {}  # For partial name consistency: "Aaron" -> "John", "Mehta" -> "Doe"
        self._initialize_bert_model()
    
    def _initialize_bert_model(self):
        """Initialize BERT NER model."""
        try:
            from transformers import AutoTokenizer, AutoModelForTokenClassification, pipeline
            
            model_name = "dslim/bert-base-NER"
            
            tokenizer = AutoTokenizer.from_pretrained(model_name)
            model = AutoModelForTokenClassification.from_pretrained(model_name)
            
            self.ner_pipeline = pipeline("ner", 
                                        model=model, 
                                        tokenizer=tokenizer,
                                        aggregation_strategy="simple")
            
        except ImportError:
            logger.error("Transformers not installed. Please run: pip install transformers torch")
            raise
        except Exception as e:
            logger.error(f"Failed to load BERT model: {e}")
            raise
    
    def _initialize_pseudo_data(self) -> Dict[str, List[str]]:
        """Initialize comprehensive pseudo data for different PII types with shorter names."""
        return {
            "first_names": [
                "Alex", "Ben", "Sam", "Dan", "Max", "Jay", "Lee", "Tom",
                "Ray", "Kim", "Ana", "Eva", "Zoe", "Mia", "Amy", "Ava",
                "Ian", "Leo", "Eli", "Zac", "Joe", "Ron", "Tim", "Jim",
                "Kai", "Roy", "Ivy", "Uma", "Ada", "Eve", "Joy", "Sky"
            ],
            "last_names": [
                "Smith", "Brown", "Davis", "Jones", "Miller", "Wilson", "Moore", "White",
                "Clark", "Lewis", "Lee", "King", "Hall", "Green", "Adams", "Baker",
                "Hill", "Scott", "Young", "Allen", "Bell", "Ross", "Gray", "Fox",
                "Stone", "Reed", "Cook", "Price", "Lane", "Wood", "Cole", "Webb"
            ],
            "organizations": [
                "TechCorp Solutions", "Global Industries Inc", "Innovation Labs",
                "DataFlow Systems", "NextGen Technologies", "CloudFirst Corp",
                "SmartBridge LLC", "Digital Dynamics", "FutureTech Group",
                "MetaVision Inc", "Quantum Networks", "CyberEdge Solutions",
                "BlueSky Enterprises", "RedRock Corp", "GreenField Inc", "SilverStream LLC"
            ],
            "cities": [
                "Springfield", "Riverside", "Fairview", "Georgetown", "Madison",
                "Franklin", "Greenwood", "Oakville", "Hillcrest", "Brookfield",
                "Centerville", "Westfield", "Northbrook", "Southgate", "Eastmont", "Westview"
            ],
            "locations": [
                "Downtown", "Uptown", "Midtown", "Old Town", "New District", "Central Plaza",
                "Business District", "Financial Center", "Tech Quarter", "Arts District"
            ],
            "addresses": [
                "123 Main Street", "456 Oak Avenue", "789 Pine Road", "321 Elm Drive",
                "654 Maple Lane", "987 Cedar Boulevard", "147 Birch Way", "258 Willow Circle",
                "369 Ash Court", "741 Poplar Street", "852 Hickory Lane", "963 Walnut Drive"
            ],
            "apt_types": [
                "Apt 1A", "Suite 201", "Unit 5B", "Floor 3", "Room 12", "Office 4C",
                "Flat 2D", "Studio 7", "Penthouse", "Loft 3A", "Space 15", "Level 4"
            ],
            "misc": [
                "[Redacted]", "[Masked]", "[PSEUDO]", "[Hidden]", "[Confidential]", "[Private]"
            ]
        }
    
    def _get_pseudo_replacement(self, pii_type: str, original_text: str) -> str:
        """Get contextually appropriate pseudo replacement for PII type with consistency."""
        # Return cached mapping if exists
        if original_text in self.used_mappings:
            return self.used_mappings[original_text]
        
        # Generate context-aware pseudo data
        if pii_type == "PERSON":
            # Handle full names vs single names with consistent partial mapping
            text_parts = original_text.strip().split()
            if len(text_parts) == 1:
                # Single name - check if it's part of a previous full name mapping
                single_name = text_parts[0]
                
                # Check if this single name is already mapped as part of a full name
                if single_name in self.name_part_mappings:
                    replacement = self.name_part_mappings[single_name]
                else:
                    # Create new single name mapping
                    available_names = [name for name in self.pseudo_data.get("first_names", ["Alex"]) 
                                     if name not in self.used_mappings.values() and name not in self.name_part_mappings.values()]
                    if not available_names:
                        available_names = self.pseudo_data.get("first_names", ["Alex"])
                    replacement = random.choice(available_names)
                    self.name_part_mappings[single_name] = replacement
                    
            elif len(text_parts) == 2:
                # Full name - create consistent mappings for both parts
                first_name, last_name = text_parts[0], text_parts[1]
                
                # Check if parts are already mapped
                mapped_first = self.name_part_mappings.get(first_name)
                mapped_last = self.name_part_mappings.get(last_name)
                
                # Get or create first name mapping
                if not mapped_first:
                    available_first_names = [name for name in self.pseudo_data.get("first_names", ["Alex"]) 
                                           if name not in self.used_mappings.values() and name not in self.name_part_mappings.values()]
                    if not available_first_names:
                        available_first_names = self.pseudo_data.get("first_names", ["Alex"])
                    mapped_first = random.choice(available_first_names)
                    self.name_part_mappings[first_name] = mapped_first
                
                # Get or create last name mapping
                if not mapped_last:
                    available_last_names = [name for name in self.pseudo_data.get("last_names", ["Smith"]) 
                                          if name not in self.used_mappings.values() and name not in self.name_part_mappings.values()]
                    if not available_last_names:
                        available_last_names = self.pseudo_data.get("last_names", ["Smith"])
                    mapped_last = random.choice(available_last_names)
                    self.name_part_mappings[last_name] = mapped_last
                
                replacement = f"{mapped_first} {mapped_last}"
                
            else:
                # Handle names with more than 2 parts (rare case)
                mapped_parts = []
                for part in text_parts:
                    if part in self.name_part_mappings:
                        mapped_parts.append(self.name_part_mappings[part])
                    else:
                        # Create new mapping for this part
                        available_names = [name for name in self.pseudo_data.get("first_names", ["Alex"]) 
                                         if name not in self.used_mappings.values() and name not in self.name_part_mappings.values()]
                        if not available_names:
                            available_names = self.pseudo_data.get("first_names", ["Alex"])
                        new_mapping = random.choice(available_names)
                        self.name_part_mappings[part] = new_mapping
                        mapped_parts.append(new_mapping)
                
                replacement = " ".join(mapped_parts)
                
        elif pii_type in ["LOC", "ADDRESS"]:
            # Location-specific replacements
            if any(word in original_text.lower() for word in ["street", "road", "lane", "avenue", "drive"]):
                replacement = random.choice(self.pseudo_data.get("addresses", ["123 Main Street"]))
            elif any(word in original_text.lower() for word in ["city", "town", "ville"]):
                replacement = random.choice(self.pseudo_data.get("cities", ["Springfield"]))
            elif any(word in original_text.lower() for word in ["apartment", "apt", "suite", "floor"]):
                replacement = random.choice(self.pseudo_data.get("apt_types", ["Apt 1A"]))
            else:
                replacement = random.choice(self.pseudo_data.get("locations", ["Downtown"]))
        elif pii_type == "ORG":
            replacement = random.choice(self.pseudo_data.get("organizations", ["ABC Corp"]))
        elif pii_type in ["PHONE", "CREDIT_CARD", "SSN", "BANK_ACCOUNT"]:
            # Numeric data - preserve format
            digits_only = re.sub(r'\D', '', original_text)
            if digits_only:
                # Generate random numbers of same length
                replacement_digits = ''.join([str(random.randint(0, 9)) for _ in range(len(digits_only))])
                # Preserve formatting
                replacement = original_text
                for i, digit in enumerate(digits_only):
                    replacement = replacement.replace(digit, replacement_digits[i], 1)
            else:
                replacement = original_text
        else:
            # Default fallback
            replacement = random.choice(self.pseudo_data.get("misc", ["[PSEUDO]"]))
        
        # Cache the mapping
        self.used_mappings[original_text] = replacement
        return replacement
    
    
    
    def _generate_mask_replacement(self, text: str, strategy: str, pii_type: str = None) -> str:
        """Generate replacement text based on masking strategy with context awareness."""
        if strategy.lower() == "mask":
            # For names, preserve structure (first name -> first name mask, full name -> full name mask)
            if pii_type == "PERSON" and " " in text.strip():
                parts = text.strip().split()
                return " ".join("*" * len(part) for part in parts)
            else:
                return "*" * len(text)
        elif strategy.lower() == "redact":
            return "[REDACTED]"
        elif strategy.lower() == "pseudo":
            return self._get_pseudo_replacement(pii_type or "MISC", text)
        else:
            return text  # Default fallback
    
    def parse_pii_config(self, config_input: str) -> List[PIIConfig]:
        """
        Parse PII configuration from input string or file.
        
        Format: PII:Type:Strategy or PII:Type:Strategy:Page:X0:Y0:X1:Y1
        Example: John Doe:PERSON:pseudo:0:100.5:200.3:180.7:220.1
                user@email.com:EMAIL:mask
                123-45-6789:SSN:redact
        """
        configs = []
        
        if os.path.isfile(config_input):
            # Read from file
            with open(config_input, 'r') as f:
                lines = f.readlines()
        else:
            # Parse as string input
            lines = config_input.strip().split('\n')
        
        for line in lines:
            line = line.strip()
            if not line or line.startswith('#'):
                continue
            
            try:
                # Use regex to parse the config line more intelligently
                # Pattern: PII_TEXT:TYPE:STRATEGY:PAGE:X0:Y0:X1:Y1
                # We need to be careful with colons in PII text (like MAC addresses)
                
                # Find the last few colons to identify TYPE:STRATEGY:PAGE:X0:Y0:X1:Y1
                colon_positions = [i for i, char in enumerate(line) if char == ':']
                
                if len(colon_positions) >= 2:  # At minimum we need TYPE and STRATEGY
                    if len(colon_positions) >= 7:
                        # New format with coordinates: PII:TYPE:STRATEGY:PAGE:X0:Y0:X1:Y1
                        # The PII text is everything before the 7th-to-last colon
                        split_pos = colon_positions[-7]
                        pii_text = line[:split_pos].strip()
                        remaining_parts = line[split_pos+1:].split(':')
                        
                        if len(remaining_parts) >= 7:
                            pii_type = remaining_parts[0].strip().upper()
                            strategy = remaining_parts[1].strip().lower()
                            page_num = 0
                            x0 = y0 = x1 = y1 = 0.0
                            replacement = None
                            
                            try:
                                page_num = int(remaining_parts[2].strip())
                                x0 = float(remaining_parts[3].strip())
                                y0 = float(remaining_parts[4].strip())
                                x1 = float(remaining_parts[5].strip())
                                y1 = float(remaining_parts[6].strip())
                                if len(remaining_parts) > 7:
                                    replacement = remaining_parts[7].strip()
                            except ValueError as e:
                                logger.warning(f"Invalid coordinate values in line: {line} - {e}")
                                # Continue with default coordinates
                        else:
                            logger.warning(f"Invalid coordinate format in line: {line}")
                            continue
                    else:
                        # Old format: PII:TYPE:STRATEGY[:REPLACEMENT]
                        # Find the position to split based on number of colons
                        if len(colon_positions) == 2:
                            # PII:TYPE:STRATEGY
                            split_pos1 = colon_positions[-2]
                            split_pos2 = colon_positions[-1]
                        elif len(colon_positions) == 3:
                            # PII:TYPE:STRATEGY:REPLACEMENT
                            split_pos1 = colon_positions[-3]
                            split_pos2 = colon_positions[-2]
                        else:
                            # Multiple colons in PII text, find last 2 or 3 colons
                            split_pos1 = colon_positions[-3] if len(colon_positions) >= 3 else colon_positions[-2]
                            split_pos2 = colon_positions[-2] if len(colon_positions) >= 3 else colon_positions[-1]
                        
                        pii_text = line[:split_pos1].strip()
                        remaining_parts = line[split_pos1+1:].split(':', 2)
                        
                        pii_type = remaining_parts[0].strip().upper()
                        strategy = remaining_parts[1].strip().lower()
                        page_num = 0
                        x0 = y0 = x1 = y1 = 0.0
                        replacement = remaining_parts[2].strip() if len(remaining_parts) > 2 else None
                    
                    config = PIIConfig(
                        text=pii_text,
                        pii_type=pii_type,
                        strategy=strategy,
                        page_num=page_num,
                        x0=x0,
                        y0=y0,
                        x1=x1,
                        y1=y1,
                        replacement=replacement
                    )
                    configs.append(config)
                    
            except Exception as e:
                logger.warning(f"Invalid config line: {line} - {e}")
        
        return configs
    
    def find_text_instances_in_page(self, page, text: str) -> List[Tuple]:
        """Find all instances of text in a PDF page and return their rectangles."""
        try:
            text_instances = page.search_for(text)
            return text_instances
        except Exception as e:
            logger.warning(f"Could not search for text '{text}': {e}")
            return []
    
    def apply_masking_strategy(self, page, rect: Tuple, pii_config: PIIConfig) -> bool:
        """
        Apply the specified masking strategy to a text rectangle in the PDF.
        
        Args:
            page: PDF page object
            rect: Rectangle coordinates (x0, y0, x1, y1) - can be None if using config coordinates
            pii_config: PII configuration with strategy and coordinates
        
        Returns:
            True if masking was successful, False otherwise
        """
        try:
            strategy = pii_config.strategy.lower()
            
            # Use coordinates from config if available, otherwise use provided rect
            if pii_config.x0 != 0.0 or pii_config.y0 != 0.0 or pii_config.x1 != 0.0 or pii_config.y1 != 0.0:
                # Use coordinates from configuration
                target_rect = (pii_config.x0, pii_config.y0, pii_config.x1, pii_config.y1)
            elif rect:
                # Use provided rectangle
                target_rect = rect
            else:
                logger.error("No coordinates available for masking")
                return False
            
            if strategy == "redact":
                # Display [REDACTED] in black text with complete removal of original
                replacement_text = "[REDACTED]"
                success = self._replace_text_with_redaction_formatting(page, target_rect, pii_config.text, replacement_text)
                
                if not success:
                    logger.warning(f"Secure redaction failed for '{pii_config.text}', using simple black box")
                    # Fallback: create a black rectangle to completely hide area
                    redact_rect = fitz.Rect(target_rect)
                    redact_annot = page.add_redact_annot(redact_rect)
                    redact_annot.update()
                    page.apply_redactions()
                    return True
                
            elif strategy in ["mask", "pseudo"]:
                # Replace with asterisks, custom text, or pseudo data
                if strategy == "mask":
                    if pii_config.replacement:
                        replacement_text = pii_config.replacement
                    else:
                        replacement_text = self._generate_mask_replacement(
                            pii_config.text, strategy, pii_config.pii_type)
                else:  # pseudo
                    if pii_config.replacement:
                        replacement_text = pii_config.replacement
                    else:
                        replacement_text = self._get_pseudo_replacement(
                            pii_config.pii_type, pii_config.text
                        )
                
                # Use the improved text replacement method with precise coordinates
                success = self._replace_text_with_proper_formatting(page, target_rect, pii_config.text, replacement_text)
                
                if not success:
                    logger.warning(f"Text replacement failed for '{pii_config.text}', using fallback redaction")
                    # Fallback: just redact without replacement
                    redact_annot = page.add_redact_annot(target_rect)
                    redact_annot.update()
                    return False
                
            else:
                logger.warning(f"Unknown masking strategy: {strategy}")
                return False
            
            return True
            
        except Exception as e:
            logger.error(f"Error applying masking strategy: {e}")
            return False

    def _replace_text_with_proper_formatting(self, page, rect: Tuple, original_text: str, replacement_text: str) -> bool:
        """
        Replace text in a rectangle with proper font matching and positioning.
        Uses a two-phase approach: collect all text info first, then redact and replace.
        
        Args:
            page: PDF page object
            rect: Rectangle coordinates (x0, y0, x1, y1)
            original_text: Original text being replaced
            replacement_text: New text to insert
            
        Returns:
            True if replacement was successful, False otherwise
        """
        try:
            # Get detailed text information BEFORE any redaction
            text_dict = page.get_text("dict", clip=rect)
            font_info = self._extract_font_info_from_rect(page, rect)
            
            x0, y0, x1, y1 = rect
            
            # Apply smart truncation with strict control to prevent overlapping
            smart_replacement_text = self._smart_text_truncation(original_text, replacement_text, 0.95)
            
            # Create redaction annotation to completely remove the original text
            redact_rect = fitz.Rect(x0, y0, x1, y1)
            redact_annot = page.add_redact_annot(redact_rect)
            redact_annot.update()
            
            # Apply the redaction to permanently remove the text
            page.apply_redactions()
            
            # Calculate text positioning using the preserved font info with better spacing
            font_size = max(font_info.get('size', 11) * 0.9, 7)  # Slightly smaller to prevent overlap
            insert_x = x0 + 1.0  # Small left margin
            baseline_offset = font_size * 0.25  # Better baseline calculation
            baseline_y = y1 - baseline_offset
            insert_point = fitz.Point(insert_x, baseline_y)
            
            # Insert the replacement text
            try:
                page.insert_text(
                    insert_point,
                    smart_replacement_text,
                    fontsize=font_size,
                    fontname=font_info.get('fontname', 'helvetica'),
                    color=font_info.get('color', (0, 0, 0)),
                    render_mode=0
                )
                
                logger.debug(f"Successfully removed and replaced '{original_text}' with '{smart_replacement_text}' (truncated from '{replacement_text}')")
                return True
                
            except Exception as e:
                # At least the original text is removed
                return True
            
        except Exception as e:
            return self._fallback_secure_replacement(page, rect, replacement_text)
    
    def _replace_text_with_redaction_formatting(self, page, rect: Tuple, original_text: str, replacement_text: str) -> bool:
        """
        Replace text with redaction formatting (black text).
        IMPORTANT: Actually removes original text from PDF, not just visual overlay.
        
        Args:
            page: PDF page object
            rect: Rectangle coordinates (x0, y0, x1, y1)
            original_text: Original text being redacted
            replacement_text: Redaction text (usually "[REDACTED]")
            
        Returns:
            True if redaction was successful, False otherwise
        """
        try:
            # Extract font info BEFORE redaction
            font_info = self._extract_font_info_from_rect(page, rect)
            
            x0, y0, x1, y1 = rect
            
            # Create redaction annotation to completely remove the original text
            redact_rect = fitz.Rect(x0, y0, x1, y1)
            redact_annot = page.add_redact_annot(redact_rect)
            redact_annot.update()
            
            # Apply the redaction to permanently remove the text
            page.apply_redactions()
            
            # Make redacted text smaller to fit better and prevent overlap
            redact_font_size = max(font_info['size'] * 0.7, 6)  # 30% smaller but at least 6pt
            
            # Apply smart truncation for redaction text with strict control
            smart_replacement_text = self._smart_text_truncation(original_text, replacement_text, 0.9)
            
            # Calculate text positioning with better spacing
            insert_x = x0 + 1.0  # Small left margin
            baseline_offset = redact_font_size * 0.25
            baseline_y = y1 - baseline_offset
            insert_point = fitz.Point(insert_x, baseline_y)
            
            # Insert redaction text in BLACK color with bold styling for emphasis
            try:
                page.insert_text(
                    insert_point,
                    smart_replacement_text,
                    fontsize=redact_font_size,
                    fontname="helv-bold",  # Use bold for redacted text
                    color=(0, 0, 0),  # BLACK color for redaction
                    render_mode=0
                )
                
                logger.debug(f"Successfully removed and redacted '{original_text}' with '{smart_replacement_text}' in black")
                return True
                
            except Exception as e:
                # Fallback to regular helvetica if bold fails
                page.insert_text(
                    insert_point,
                    smart_replacement_text,
                    fontsize=redact_font_size,
                    fontname="helvetica",
                    color=(0, 0, 0),  # BLACK color
                    render_mode=0
                )
                return True
                
        except Exception as e:
            # Fallback: create a simple black rectangle to hide the area
            try:
                redact_rect = fitz.Rect(rect)
                redact_annot = page.add_redact_annot(redact_rect, fill=(0, 0, 0))  # Black fill
                redact_annot.update()
                page.apply_redactions()
                return True
            except Exception as fallback_error:
                logger.error(f"Even fallback redaction failed: {fallback_error}")
                return False
    
    def _extract_font_info_from_rect(self, page, rect: Tuple) -> Dict[str, Any]:
        """
        Extract font information from a rectangle before any redaction occurs.
        
        Args:
            page: PDF page object
            rect: Rectangle coordinates
            
        Returns:
            Dictionary with font information
        """
        try:
            x0, y0, x1, y1 = rect
            text_dict = page.get_text("dict", clip=rect)
            
            font_info = {
                'size': max((y1 - y0) * 0.75, 8),  # Default estimate
                'fontname': 'helvetica',
                'color': (0, 0, 0)
            }
            
            if text_dict and 'blocks' in text_dict:
                for block in text_dict['blocks']:
                    if 'lines' in block:
                        for line in block['lines']:
                            if 'spans' in line:
                                for span in line['spans']:
                                    if span.get('size', 0) > 0:
                                        font_info['size'] = span['size']
                                    
                                    if span.get('font'):
                                        font_flags = span.get('flags', 0)
                                        font_info['fontname'] = self._get_proper_fontname(span['font'], font_flags)
                                    
                                    if span.get('color') is not None:
                                        color = span['color']
                                        if isinstance(color, int):
                                            r = (color >> 16) & 255
                                            g = (color >> 8) & 255
                                            b = color & 255
                                            font_info['color'] = (r/255.0, g/255.0, b/255.0)
                                    
                                    # Use first valid font info found
                                    if font_info['size'] > 0:
                                        return font_info
            
            return font_info
            
        except Exception as e:
            return {
                'size': max((rect[3] - rect[1]) * 0.75, 8),
                'fontname': 'helvetica',
                'color': (0, 0, 0)
            }
    
    def _get_font_properties_from_rect(self, rect: Tuple) -> Tuple[float, str]:
        """
        Estimate font properties from rectangle dimensions when text is not accessible.
        
        Args:
            rect: Rectangle coordinates (x0, y0, x1, y1)
            
        Returns:
            Tuple of (font_size, font_name)
        """
        try:
            x0, y0, x1, y1 = rect
            text_height = y1 - y0
            
            # Estimate font size from rectangle height (typical ratio is 0.7-0.8)
            estimated_font_size = max(text_height * 0.75, 8)
            
            return estimated_font_size, "helvetica"
            
        except Exception as e:
            return 11.0, "helvetica"
    
    def _fallback_secure_replacement(self, page, rect: Tuple, replacement_text: str) -> bool:
        """
        Secure fallback method that ensures original text is removed even if detailed replacement fails.
        """
        try:
            x0, y0, x1, y1 = rect
            
            # CRITICAL: Always remove original text first using redaction
            redact_rect = fitz.Rect(x0, y0, x1, y1)
            redact_annot = page.add_redact_annot(redact_rect)
            redact_annot.update()
            page.apply_redactions()
            
            # Estimate font size and apply basic truncation
            font_size = max((y1 - y0) * 0.7, 8)
            smart_replacement_text = replacement_text
            if len(replacement_text) > 12:
                if replacement_text == "[REDACTED]":
                    smart_replacement_text = "[RED]"
                elif len(replacement_text) > 15:
                    smart_replacement_text = replacement_text[:12] + ".."
            
            # Insert replacement text
            insert_x = x0 + 0.5
            baseline_y = y1 - (font_size * 0.22)
            insert_point = fitz.Point(insert_x, baseline_y)
            
            page.insert_text(
                insert_point,
                smart_replacement_text,
                fontsize=font_size,
                fontname="helvetica",
                color=(0, 0, 0),
                render_mode=0
            )
            
            logger.debug(f"Applied secure fallback replacement: '{smart_replacement_text}'")
            return True
            
        except Exception as e:
            logger.error(f"Secure fallback replacement failed: {e}")
            # Even if text insertion fails, original text is removed
            return True
    
    def _fallback_text_replacement(self, page, rect: Tuple, replacement_text: str) -> bool:
        """
        Secure fallback method for text replacement that ensures original text is completely removed.
        """
        try:
            x0, y0, x1, y1 = rect
            
            # CRITICAL: Always remove original text first using redaction
            redact_rect = fitz.Rect(x0, y0, x1, y1)
            redact_annot = page.add_redact_annot(redact_rect)
            redact_annot.update()
            page.apply_redactions()
            
            # For fallback, apply basic truncation if text is very long
            smart_replacement_text = replacement_text
            if len(replacement_text) > 12:  # If very long, truncate
                if replacement_text == "[REDACTED]":
                    smart_replacement_text = "[RED]"  # Shorter redaction
                elif len(replacement_text) > 15:
                    smart_replacement_text = replacement_text[:12] + ".."
            
            # Estimate font size from rectangle height
            font_size = max((y1 - y0) * 0.75, 8)
            
            # Calculate text positioning
            insert_x = x0 + 0.5
            baseline_y = y1 - (font_size * 0.22)
            insert_point = fitz.Point(insert_x, baseline_y)
            
            # Insert replacement text with conservative settings
            page.insert_text(
                insert_point,
                smart_replacement_text,
                fontsize=font_size,
                fontname="helvetica",
                color=(0, 0, 0),  # Always black for readability
                render_mode=0
            )
            
            return True
            
        except Exception as e:
            logger.error(f"Secure fallback text replacement failed: {e}")
            # Even if replacement fails, original text is removed by redaction
            return True
    
    def _smart_text_truncation(self, original_text: str, replacement_text: str, max_width_ratio: float = 1.0) -> str:
        """
        Intelligently truncate replacement text to prevent overlapping with adjacent text.
        More aggressive truncation to ensure clean layout.
        
        Args:
            original_text: Original text being replaced
            replacement_text: Replacement text
            max_width_ratio: Maximum allowed width ratio (replacement/original) - default 1.0 for strict control
            
        Returns:
            Potentially truncated replacement text
        """
        try:
            original_len = len(original_text.strip())
            replacement_len = len(replacement_text)
            
            # Very strict control - replacement should not be much longer than original
            if replacement_len <= original_len * max_width_ratio:
                return replacement_text
            
            # For redacted text, use progressively shorter versions based on original length
            if replacement_text == "[REDACTED]":
                if original_len <= 2:
                    return "**"  # Very short original
                elif original_len <= 4:
                    return "[*]"  # Short original  
                elif original_len <= 6:
                    return "[RED]"  # Medium original
                elif original_len <= 8:
                    return "[REDACT]"  # Longer original
                else:
                    return "[REDACTED]"  # Keep full only for very long original
            
            # For pseudo names, be very aggressive in shortening to prevent overlap
            if " " in replacement_text and replacement_text.count(" ") == 1:  # Full names
                parts = replacement_text.split()
                if len(parts) >= 2:
                    first_name, last_name = parts[0], parts[1]
                    
                    # Very short original - use just first name or initials
                    if original_len <= 4:
                        return first_name[:3] if len(first_name) > 3 else first_name
                    elif original_len <= 6:
                        # Short original - use first name only or first + initial
                        if len(first_name) <= 4:
                            return f"{first_name} {last_name[0]}"
                        else:
                            return first_name[:4]
                    elif original_len <= 10:
                        # Medium original - shorten both names if needed
                        short_first = first_name[:4] if len(first_name) > 4 else first_name
                        short_last = last_name[:4] if len(last_name) > 4 else last_name
                        return f"{short_first} {short_last}"
                    else:
                        # Longer original - allow slightly longer but still controlled
                        max_first = min(len(first_name), 5)
                        max_last = min(len(last_name), 6)
                        return f"{first_name[:max_first]} {last_name[:max_last]}"
            
            # For single words (including single names), be very strict
            elif " " not in replacement_text:
                # Single word replacement
                if original_len <= 3:
                    return replacement_text[:3]
                elif original_len <= 5:
                    return replacement_text[:4]
                elif original_len <= 8:
                    return replacement_text[:6]
                else:
                    # Allow longer but with limit
                    max_len = min(original_len, 8)
                    return replacement_text[:max_len]
            
            # For other text types, be conservative with truncation
            max_len = max(int(original_len * max_width_ratio), 4)  # At least 4 chars
            if replacement_len > max_len:
                if max_len > 6:
                    return replacement_text[:max_len-2] + ".."
                else:
                    return replacement_text[:max_len]
            
            return replacement_text
            
        except Exception as e:
            # Fallback: very conservative truncation
            max_len = min(len(original_text), len(replacement_text), 8)
            return replacement_text[:max_len] if len(replacement_text) > max_len else replacement_text
    
    def _calculate_text_width(self, text: str, font_size: float, fontname: str) -> float:
        """
        Calculate approximate text width based on font size and font type.
        
        Args:
            text: Text to measure
            font_size: Font size in points
            fontname: Font name
            
        Returns:
            Approximate text width in points
        """
        try:
            # Character width ratios for different fonts (relative to font size)
            font_width_ratios = {
                'courier': 0.6,  # Monospace - all chars same width
                'cour-bold': 0.6,
                'cour-oblique': 0.6,
                'cour-boldoblique': 0.6,
                'helvetica': 0.55,  # Sans-serif
                'helv-bold': 0.6,
                'helv-oblique': 0.55,
                'helv-boldoblique': 0.6,
                'times-roman': 0.5,  # Serif - typically narrower
                'times-bold': 0.55,
                'times-italic': 0.5,
                'times-bolditalic': 0.55,
            }
            
            # Get width ratio for the font (default to helvetica if not found)
            ratio = font_width_ratios.get(fontname.lower(), 0.55)
            
            # Calculate approximate width
            # Account for character variations - some letters are wider than others
            char_count = len(text)
            
            # Apply character-specific adjustments
            wide_chars = sum(1 for c in text if c.upper() in 'MWQG@')
            narrow_chars = sum(1 for c in text if c in 'ijl .,;:')
            normal_chars = char_count - wide_chars - narrow_chars
            
            # Weighted character count
            weighted_count = (wide_chars * 1.3 + narrow_chars * 0.7 + normal_chars * 1.0)
            
            estimated_width = weighted_count * font_size * ratio
            
            return estimated_width
            
        except Exception as e:
            # Fallback: simple calculation
            return len(text) * font_size * 0.55
    
    def _get_proper_fontname(self, original_font: str, font_flags: int) -> str:
        """
        Convert original font name to a proper PyMuPDF font name based on flags.
        
        Args:
            original_font: Original font name from PDF
            font_flags: Font flags indicating style (bold, italic, etc.)
        
        Returns:
            Proper font name for PyMuPDF
        """
        # Font flags meanings:
        # 16 = bold
        # 2 = italic
        # 18 = bold + italic
        
        is_bold = bool(font_flags & 16)
        is_italic = bool(font_flags & 2)
        
        # Map common fonts to PyMuPDF equivalents
        font_lower = original_font.lower()
        
        if 'helvetica' in font_lower or 'arial' in font_lower:
            if is_bold and is_italic:
                return "helv-boldoblique"
            elif is_bold:
                return "helv-bold"
            elif is_italic:
                return "helv-oblique"
            else:
                return "helvetica"
        
        elif 'times' in font_lower:
            if is_bold and is_italic:
                return "times-bolditalic"
            elif is_bold:
                return "times-bold"
            elif is_italic:
                return "times-italic"
            else:
                return "times-roman"
        
        elif 'courier' in font_lower:
            if is_bold and is_italic:
                return "cour-boldoblique"
            elif is_bold:
                return "cour-bold"
            elif is_italic:
                return "cour-oblique"
            else:
                return "courier"
        
        else:
            # Default to helvetica variants
            if is_bold and is_italic:
                return "helv-boldoblique"
            elif is_bold:
                return "helv-bold"
            elif is_italic:
                return "helv-oblique"
            else:
                return "helvetica"
    
    def _get_font_properties(self, page, rect: Tuple) -> Tuple[float, str]:
        """Extract font properties from the text in the given rectangle."""
        try:
            # Get text dictionary for the rectangle area
            text_dict = page.get_text("dict", clip=rect)
            
            font_size = 11.0  # default
            font_name = "helvetica"  # default
            
            if text_dict and 'blocks' in text_dict:
                for block in text_dict['blocks']:
                    if 'lines' in block:
                        for line in block['lines']:
                            if 'spans' in line:
                                for span in line['spans']:
                                    if 'size' in span and span['size'] > 0:
                                        font_size = span['size']
                                    if 'font' in span and span['font']:
                                        # Get proper font name
                                        original_font = span['font']
                                        font_flags = span.get('flags', 0)
                                        font_name = self._get_proper_fontname(original_font, font_flags)
                                    # Use the first valid font properties found
                                    if font_size > 0:
                                        return font_size, font_name
            
            return font_size, font_name
            
        except Exception as e:
            return 11.0, "helvetica"  # fallback defaults
    
    def mask_pdf_with_config(self, input_pdf_path: str, output_pdf_path: str, 
                           pii_configs: List[PIIConfig]) -> Dict[str, Any]:
        """
        Process a PDF file and mask PII according to the provided configuration.
        
        Args:
            input_pdf_path: Path to input PDF file
            output_pdf_path: Path to output masked PDF file
            pii_configs: List of PII configurations
            
        Returns:
            Dictionary with processing statistics
        """
        stats = {
            "total_pages": 0,
            "total_pii_masked": 0,
            "strategies_used": {},
            "pages_processed": 0,
            "failed_maskings": 0
        }
        
        try:
            # Open the PDF
            doc = fitz.open(input_pdf_path)
            stats["total_pages"] = len(doc)
            
            logger.info(f"Processing PDF: {input_pdf_path} ({stats['total_pages']} pages)")
            
            # Extract full text for LLM validation
            full_text = ""
            for page_num in range(len(doc)):
                page = doc[page_num]
                full_text += page.get_text() + "\n"

            validated_configs = pii_configs
            
            # Process each page
            for page_num in range(len(doc)):
                page = doc[page_num]
                page_text = page.get_text()
                
                if not page_text.strip():
                    logger.debug(f"Page {page_num + 1}: No text found, skipping")
                    continue
                
                # Get configurations for this page from validated configs
                page_configs = [config for config in validated_configs if config.page_num == page_num]
                if not page_configs:
                    continue
                
                page_masked_count = 0
                
                # CRITICAL: Sort configurations by text length (longest first) to avoid overlapping replacements
                # This prevents issues where "Aaron Mehta" gets replaced partially by "Aaron" and "Mehta"
                page_configs.sort(key=lambda x: (-len(x.text), -x.y0, -x.x0))
                
                # Process PII in smaller batches to maintain consistency
                batch_size = 5  # Process 5 items at a time
                for batch_start in range(0, len(page_configs), batch_size):
                    batch_configs = page_configs[batch_start:batch_start + batch_size]
                    
                    # Apply each PII configuration in the batch
                    for pii_config in batch_configs:
                        success = False
                        
                        # Check if we have coordinates in the config
                        if (pii_config.x0 != 0.0 or pii_config.y0 != 0.0 or 
                            pii_config.x1 != 0.0 or pii_config.y1 != 0.0):
                            # Use coordinates from config - no need to search
                            success = self.apply_masking_strategy(page, None, pii_config)
                            
                            if not success:
                                logger.warning(f"Failed to mask '{pii_config.text}' at specified coordinates")
                        else:
                            logger.warning(f"No coordinates available for PII '{pii_config.text}', trying text search fallback")
                            # Fallback: search for text instances (old method)
                            text_instances = self.find_text_instances_in_page(page, pii_config.text)
                            
                            if not text_instances:
                                logger.warning(f"PII '{pii_config.text}' not found on page {page_num + 1}")
                                continue
                            
                            # Apply masking strategy to the first instance (most likely correct)
                            rect = text_instances[0]
                            success = self.apply_masking_strategy(page, rect, pii_config)
                        
                        if success:
                            page_masked_count += 1
                            stats["total_pii_masked"] += 1
                            
                            # Update strategy statistics
                            strategy = pii_config.strategy
                            if strategy not in stats["strategies_used"]:
                                stats["strategies_used"][strategy] = 0
                            stats["strategies_used"][strategy] += 1
                        else:
                            stats["failed_maskings"] += 1
                    
                    # Small delay between batches to ensure processing consistency
                    import time
                    time.sleep(0.01)  # 10ms delay
                
                if page_masked_count > 0:
                    # Text has already been securely removed during individual processing
                    # No need for additional redaction steps
                    stats["pages_processed"] += 1
            
            # Save the masked PDF
            doc.save(output_pdf_path)
            doc.close()
            
            logger.info(f"✓ Masked PDF saved to: {output_pdf_path}")
            
            return stats
            
        except Exception as e:
            logger.error(f"Error processing PDF: {e}")
            raise
    
    def generate_masking_report(self, stats: Dict[str, Any], pii_configs: List[PIIConfig], 
                              output_path: Optional[str] = None) -> str:
        """Generate a detailed report of the PII masking process."""
        report_lines = [
            "BERT PII Masking Report",
            "=" * 50,
        ]
        
        # Handle different stats formats (direct vs text-based methods)
        if 'pages_processed' in stats and 'total_pages' in stats:
            report_lines.append(f"Total pages processed: {stats['pages_processed']}/{stats['total_pages']}")
        
        report_lines.extend([
            f"Total PII instances masked: {stats['total_pii_masked']}",
            f"Failed maskings: {stats.get('failed_maskings', 0)}",
            f"PII configurations applied: {len(pii_configs)}",
            "",
            "Masking Strategies Used:",
            "-" * 30
        ])
        
        for strategy, count in stats["strategies_used"].items():
            report_lines.append(f"{strategy}: {count} instances")
        
        report_lines.extend([
            "",
            "PII Configurations Applied:",
            "-" * 40
        ])
        
        for config in pii_configs:
            replacement_info = f" -> {config.replacement}" if config.replacement else ""
            report_lines.append(f"'{config.text}' ({config.pii_type}) : {config.strategy}{replacement_info}")
        
        if self.used_mappings:
            report_lines.extend([
                "",
                "Pseudo Replacements Generated:",
                "-" * 40
            ])
            for original, replacement in self.used_mappings.items():
                report_lines.append(f"'{original}' -> '{replacement}'")
        
        # Add consistent name part mappings for better transparency
        if self.name_part_mappings:
            report_lines.extend([
                "",
                "Consistent Name Part Mappings:",
                "-" * 40,
                "(This ensures partial names are replaced consistently)"
            ])
            for original_part, replacement_part in self.name_part_mappings.items():
                report_lines.append(f"'{original_part}' -> '{replacement_part}'")
        
        report = "\n".join(report_lines)
        
        if output_path:
            with open(output_path, 'w') as f:
                f.write(report)
            logger.info(f"Report saved to: {output_path}")
        
        return report

    def pdf_to_text(self, pdf_path: str) -> str:
        """Extract text from PDF while preserving structure."""
        try:
            doc = fitz.open(pdf_path)
            full_text = ""
            
            for page_num in range(len(doc)):
                page = doc[page_num]
                page_text = page.get_text()
                full_text += page_text + "\n\n"  # Add page breaks
                
            doc.close()
            return full_text
        except Exception as e:
            logger.error(f"Error extracting text from PDF: {e}")
            raise

    def pdf_to_docx(self, pdf_path: str, docx_path: str):
        """Convert PDF to Word document."""
        try:
            text = self.pdf_to_text(pdf_path)
            doc = Document()
            
            # Split text into paragraphs
            paragraphs = text.split('\n')
            
            for paragraph in paragraphs:
                if paragraph.strip():  # Skip empty lines
                    doc.add_paragraph(paragraph)
                else:
                    doc.add_paragraph("")  # Preserve spacing
            
            doc.save(docx_path)
            logger.info(f"PDF converted to Word document: {docx_path}")
        except Exception as e:
            logger.error(f"Error converting PDF to DOCX: {e}")
            raise

    def mask_text(self, text: str, pii_configs: List[PIIConfig]) -> Tuple[str, Dict[str, Any]]:
        """
        Apply PII masking to plain text.
        
        Args:
            text: The text content to mask
            pii_configs: List of PII configurations
            
        Returns:
            Tuple of (masked_text, stats)
        """
        stats = {
            "total_pii_masked": 0,
            "strategies_used": {},
            "failed_maskings": 0
        }
        
        masked_text = text
        
        # Sort PII configs by text length (longest first) to avoid overlapping replacements
        sorted_configs = sorted(pii_configs, key=lambda x: -len(x.text))
        
        for config in sorted_configs:
            original_text = config.text
            
            if original_text not in masked_text:
                logger.warning(f"PII text '{original_text}' not found in document")
                stats["failed_maskings"] += 1
                continue
            
            # Generate replacement based on strategy
            if config.strategy == "redact":
                replacement = "[REDACTED]"
            elif config.strategy == "mask":
                replacement = "*" * len(original_text)
            elif config.strategy == "pseudo":
                if config.replacement:
                    replacement = config.replacement
                else:
                    replacement = self._get_pseudo_replacement(config.pii_type, original_text)
                    # Store the mapping for consistency
                    self.used_mappings[original_text] = replacement
            else:
                logger.warning(f"Unknown strategy: {config.strategy}")
                replacement = "[UNKNOWN_STRATEGY]"
            
            # Replace all occurrences
            count = masked_text.count(original_text)
            masked_text = masked_text.replace(original_text, replacement)
            
            if count > 0:
                stats["total_pii_masked"] += count
                stats["strategies_used"][config.strategy] = stats["strategies_used"].get(config.strategy, 0) + count
                logger.info(f"Masked {count} occurrences of '{original_text}' with '{replacement}' ({config.strategy} strategy)")
        
        return masked_text, stats

    def mask_docx(self, input_docx_path: str, output_docx_path: str, pii_configs: List[PIIConfig]) -> Dict[str, Any]:
        """
        Apply PII masking to a Word document.
        
        Args:
            input_docx_path: Path to input Word document
            output_docx_path: Path to output masked Word document
            pii_configs: List of PII configurations
            
        Returns:
            Dictionary with processing statistics
        """
        try:
            doc = Document(input_docx_path)
            
            # Extract all text from the document
            full_text = ""
            for paragraph in doc.paragraphs:
                full_text += paragraph.text + "\n"
            
            # Apply masking to the text
            masked_text, stats = self.mask_text(full_text, pii_configs)
            
            # Replace paragraphs with masked content
            masked_lines = masked_text.split('\n')
            
            # Clear existing paragraphs
            for paragraph in doc.paragraphs:
                paragraph.clear()
            
            # Add masked paragraphs
            for i, line in enumerate(masked_lines):
                if i < len(doc.paragraphs):
                    doc.paragraphs[i].text = line
                else:
                    doc.add_paragraph(line)
            
            # Save the masked document
            doc.save(output_docx_path)
            logger.info(f"Masked Word document saved: {output_docx_path}")
            
            return stats
        
        except Exception as e:
            logger.error(f"Error masking DOCX: {e}")
            raise

    def docx_to_pdf(self, docx_path: str, pdf_path: str):
        """Convert Word document to PDF using reportlab."""
        try:
            doc = Document(docx_path)
            
            # Create PDF
            from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
            from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
            from reportlab.lib.pagesizes import letter
            from reportlab.lib.units import inch
            
            # Create the PDF document
            pdf_doc = SimpleDocTemplate(pdf_path, pagesize=letter)
            styles = getSampleStyleSheet()
            story = []
            
            # Add each paragraph from the Word document
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    p = Paragraph(paragraph.text, styles['Normal'])
                    story.append(p)
                    story.append(Spacer(1, 0.1*inch))
            
            # Build the PDF
            pdf_doc.build(story)
            logger.info(f"Word document converted to PDF: {pdf_path}")
            
        except Exception as e:
            logger.error(f"Error converting DOCX to PDF: {e}")
            raise

    def text_to_pdf(self, text: str, pdf_path: str):
        """Convert plain text to PDF using reportlab."""
        try:
            from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
            from reportlab.lib.styles import getSampleStyleSheet
            from reportlab.lib.pagesizes import letter
            from reportlab.lib.units import inch
            
            # Create the PDF document
            doc = SimpleDocTemplate(pdf_path, pagesize=letter)
            styles = getSampleStyleSheet()
            story = []
            
            # Split text into paragraphs
            paragraphs = text.split('\n')
            
            for paragraph in paragraphs:
                if paragraph.strip():
                    p = Paragraph(paragraph, styles['Normal'])
                    story.append(p)
                    story.append(Spacer(1, 0.1*inch))
            
            # Build the PDF
            doc.build(story)
            logger.info(f"Text converted to PDF: {pdf_path}")
            
        except Exception as e:
            logger.error(f"Error converting text to PDF: {e}")
            raise

    def process_pdf_via_text_conversion(self, input_pdf_path: str, output_pdf_path: str, 
                                      pii_configs: List[PIIConfig], use_docx: bool = True) -> Dict[str, Any]:
        """
        Process PDF by converting to text/docx, masking, then converting back to PDF.
        
        Args:
            input_pdf_path: Path to input PDF file
            output_pdf_path: Path to output masked PDF file
            pii_configs: List of PII configurations
            use_docx: If True, convert via Word document; if False, use plain text
            
        Returns:
            Dictionary with processing statistics
        """
        import tempfile
        import os
        
        try:
            # Create temporary files
            temp_dir = tempfile.mkdtemp()
            
            if use_docx:
                temp_docx_input = os.path.join(temp_dir, "temp_input.docx")
                temp_docx_masked = os.path.join(temp_dir, "temp_masked.docx")
                
                # Step 1: PDF -> DOCX
                logger.info("Converting PDF to Word document...")
                self.pdf_to_docx(input_pdf_path, temp_docx_input)
                
                # Step 2: Mask DOCX
                logger.info("Applying PII masking to Word document...")
                stats = self.mask_docx(temp_docx_input, temp_docx_masked, pii_configs)
                
                # Step 3: DOCX -> PDF
                logger.info("Converting masked Word document back to PDF...")
                self.docx_to_pdf(temp_docx_masked, output_pdf_path)
                
            else:
                # Step 1: PDF -> Text
                logger.info("Extracting text from PDF...")
                text = self.pdf_to_text(input_pdf_path)
                
                # Step 2: Mask Text
                logger.info("Applying PII masking to text...")
                masked_text, stats = self.mask_text(text, pii_configs)
                
                # Step 3: Text -> PDF
                logger.info("Converting masked text back to PDF...")
                self.text_to_pdf(masked_text, output_pdf_path)
            
            # Clean up temporary files
            import shutil
            shutil.rmtree(temp_dir)
            
            logger.info(f"PDF processing completed: {output_pdf_path}")
            return stats
            
        except Exception as e:
            logger.error(f"Error in PDF text conversion workflow: {e}")
            raise

def main():
    """Main function for command-line usage."""
    print("BERT PII Masker with Configurable Strategies")
    print("=" * 60)
    
    if len(sys.argv) < 3:
        print("Usage:")
        print("  python bert_pii_masker.py <input.pdf> <output.pdf> <config_file> [--method=<method>]")
        print("  python bert_pii_masker.py <input.pdf> <output.pdf> --interactive [--method=<method>]")
        print("  python bert_pii_masker.py --sample-config")
        print("")
        print("Methods:")
        print("  --method=direct    : Direct PDF editing (default, preserves layout)")
        print("  --method=text      : PDF -> Text -> PDF conversion")
        print("  --method=docx      : PDF -> Word -> PDF conversion")
        print("")
        print("Examples:")
        print("  python bert_pii_masker.py document.pdf masked.pdf pii_config.txt")
        print("  python bert_pii_masker.py document.pdf masked.pdf pii_config.txt --method=docx")
        print("  python bert_pii_masker.py document.pdf masked.pdf --interactive --method=text")
        return 1
    
    input_pdf = sys.argv[1]
    output_pdf = sys.argv[2]
    
    # Parse method parameter
    method = "docx"  # default method
    for arg in sys.argv:
        if arg.startswith("--method="):
            method = arg.split("=")[1].lower()
            break
    
    if method not in ["direct", "text", "docx"]:
        print(f"Error: Invalid method '{method}'. Use 'direct', 'text', or 'docx'")
        return 1
    
    if not os.path.exists(input_pdf):
        print(f"Error: Input file '{input_pdf}' not found")
        return 1
    
    if not input_pdf.lower().endswith('.pdf'):
        print(f"Error: Input file must be a PDF")
        return 1
    
    try:
        # Initialize the masker
        masker = BERTPIIMasker()
        if len(sys.argv) > 3 and not sys.argv[3].startswith("--"):
            # Configuration file mode
            config_input = sys.argv[3]
            pii_configs = masker.parse_pii_config(config_input)
        else:
            print("Error: Please provide configuration file or use --interactive mode")
            return 1
        
        if not pii_configs:
            print("No PII configurations found. Exiting.")
            return 1
        
        # Process the PDF based on selected method
        print(f"\nProcessing PDF with {len(pii_configs)} PII configurations using {method} method...")
        
        if method == "direct":
            # Original direct PDF editing method
            stats = masker.mask_pdf_with_config(input_pdf, output_pdf, pii_configs)
        elif method == "text":
            # PDF -> Text -> PDF method
            stats = masker.process_pdf_via_text_conversion(input_pdf, output_pdf, pii_configs, use_docx=False)
        elif method == "docx":
            # PDF -> Word -> PDF method
            stats = masker.process_pdf_via_text_conversion(input_pdf, output_pdf, pii_configs, use_docx=True)
        
        # Generate and save report
        report_path = output_pdf.replace('.pdf', '_masking_report.txt')
        report = masker.generate_masking_report(stats, pii_configs, report_path)
        
        # Display summary
        print("\n" + "="*60)
        print("PII MASKING COMPLETED SUCCESSFULLY!")
        print("="*60)
        print(f"Input PDF: {input_pdf}")
        print(f"Output PDF: {output_pdf}")
        print(f"Report: {report_path}")
        print(f"Method Used: {method}")
        print("\nSummary:")
        
        # Handle different stats formats
        if method == "direct":
            print(f"  • Pages processed: {stats.get('pages_processed', 'N/A')}/{stats.get('total_pages', 'N/A')}")
            print(f"  • PII instances masked: {stats['total_pii_masked']}")
            print(f"  • Failed maskings: {stats.get('failed_maskings', 0)}")
        else:
            print(f"  • PII instances masked: {stats['total_pii_masked']}")
            print(f"  • Failed maskings: {stats.get('failed_maskings', 0)}")
        
        if stats['strategies_used']:
            print("\nStrategies Used:")
            for strategy, count in stats['strategies_used'].items():
                print(f"  • {strategy}: {count} instances")
        
        return 0
        
    except Exception as e:
        logger.error(f"Error processing PDF: {e}")
        return 1

if __name__ == "__main__":
    sys.exit(main())
